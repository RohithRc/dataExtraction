from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, ListFlowable, ListItem, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt, RGBColor
from fastapi.responses import FileResponse
from reportlab.lib.pagesizes import A4, landscape

def generate_pdf(title, metadata, headers, rows):
    filename = "cost-report.pdf"
    doc = SimpleDocTemplate(
        filename, 
        pagesize=landscape(A4),
        rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=18
    )
    styles = getSampleStyleSheet()
    
    # Custom Styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=20,
        alignment=1 # Center
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=18,
        textColor=colors.HexColor('#34495e'),
        spaceBefore=15,
        spaceAfter=10
    )
    
    elements = []

    # Metadata section
    for line in metadata:
        elements.append(Paragraph(line, styles['Normal']))
    
    elements.append(Spacer(1, 20))
    elements.append(Paragraph(title, title_style))
    
    # Prepare table data
    data = [headers] + rows
    t = Table(data)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ecf0f1')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(t)

    doc.build(elements)

    return FileResponse(
        filename,
        media_type="application/pdf",
        filename=filename
    )

def generate_docx(title, metadata, headers, rows):
    filename = "cost-report.docx"
    doc = Document()

    for line in metadata:
        doc.add_paragraph(line)
        
    doc.add_heading(title, level=1)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row style
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(11)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    doc.save(filename)

    return FileResponse(
        filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

def create_pdf_from_elements(elements):
    filename = "converted-report.pdf"
    doc = SimpleDocTemplate(
        filename, 
        pagesize=landscape(A4),
        rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50
    )
    styles = getSampleStyleSheet()
    
    # Custom Styles
    custom_h1 = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=20, textColor=colors.HexColor('#2c3e50'), spaceAfter=14)
    custom_h2 = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=16, textColor=colors.HexColor('#34495e'), spaceBefore=12, spaceAfter=10)
    custom_h3 = ParagraphStyle('H3', parent=styles['Heading3'], fontSize=14, textColor=colors.HexColor('#34495e'), spaceBefore=10, spaceAfter=8)
    
    custom_normal = ParagraphStyle('CustomNormal', parent=styles['Normal'], fontSize=11, leading=14, spaceAfter=6)
    
    custom_code = ParagraphStyle('Code', parent=styles['Code'], fontSize=9, leading=11, backColor=colors.whitesmoke, borderColor=colors.grey, borderWidth=0.5, leftIndent=10, rightIndent=10, spaceAfter=10, splitLongWords=1, wordWrap='CJK')
    
    custom_blockquote = ParagraphStyle('Blockquote', parent=styles['Normal'], leftIndent=20, rightIndent=20, textColor=colors.HexColor('#555555'), fontName='Helvetica-Oblique')

    story = []

    for el in elements:
        if el['type'] == 'heading':
            level = el['level']
            if level == 1:
                style = custom_h1
            elif level == 2:
                style = custom_h2
            else:
                style = custom_h3
            story.append(Paragraph(el['text'], style))
            
        elif el['type'] == 'paragraph':
            story.append(Paragraph(el['text'], custom_normal))
            story.append(Spacer(1, 6))
            
        elif el['type'] == 'list':
            list_items = [ListItem(Paragraph(item, custom_normal)) for item in el['items']]
            story.append(ListFlowable(
                list_items,
                bulletType='bullet' if el['style'] == 'unordered' else '1',
                start='circle',
                leftIndent=20,
                spaceAfter=10
            ))
            
        elif el['type'] == 'code':
            story.append(Preformatted(el['text'], custom_code))
            story.append(Spacer(1, 10))
            
        elif el['type'] == 'blockquote':
            story.append(Paragraph(el['text'], custom_blockquote))
            story.append(Spacer(1, 10))
            
        elif el['type'] == 'table':
            t = Table(el['data'])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#ecf0f1')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(t)
            story.append(Spacer(1, 16))
            
    doc.build(story)
    return FileResponse(filename, media_type="application/pdf", filename=filename)

def create_docx_from_elements(elements):
    filename = "converted-report.docx"
    doc = Document()

    for el in elements:
        if el['type'] == 'heading':
            doc.add_heading(el['text'], level=el['level'])
        
        elif el['type'] == 'paragraph':
            doc.add_paragraph(el['text'])
            
        elif el['type'] == 'list':
            style = 'List Bullet' if el['style'] == 'unordered' else 'List Number'
            for item in el['items']:
                try:
                    doc.add_paragraph(item, style=style)
                except:
                    # Fallback if style doesn't exist
                    p = doc.add_paragraph(item)
                    p.style = 'List Paragraph'

        elif el['type'] == 'blockquote':
            try:
                doc.add_paragraph(el['text'], style='Quote')
            except:
                p = doc.add_paragraph(el['text'])
                p.italic = True
                
        elif el['type'] == 'code':
            try:
                # Use 'Macro Text' or similar monospaced style if available
                doc.add_paragraph(el['text'], style='Macro Text') 
            except:
                 # Fallback: simple monospace
                 p = doc.add_paragraph(el['text'])
                 for run in p.runs:
                     run.font.name = 'Courier New'

        elif el['type'] == 'table':
            rows = el['data']
            if not rows: continue
            
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            
            # Header styling
            for j, cell in enumerate(table.rows[0].cells):
                cell.text = str(rows[0][j])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)

            # Data rows
            for i, row in enumerate(rows[1:]): # Skip header
                # Note: add_table creates necessary rows, but accessing them by index requires existing row objects
                # python-docx tables are fixed size upon creation unless add_row is called.
                # Here we created strict size, so we access directly.
                docx_row = table.rows[i+1]
                for j, cell_text in enumerate(row):
                    docx_row.cells[j].text = str(cell_text)

    doc.save(filename)
    return FileResponse(filename, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)
